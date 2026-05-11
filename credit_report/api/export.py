"""DOCX and PDF export endpoints for generated credit report sections."""
from __future__ import annotations

import asyncio
import html as html_mod
import io
import logging
import re
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Response
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from credit_report.database import get_db
from credit_report.models import Report, SectionOutput
from credit_report.security.auth import get_current_user
from credit_report.security.models import User

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/reports/{report_id}", tags=["export"])

SECTION_NAMES = {
    1: "Facility Structure & Key Terms",
    2: "Overall Comments",
    3: "Credit Risk Assessment",
    4: "Borrower Background",
    5: "Collateral Assessment",
    6: "Project Overview (Ship Finance)",
    7: "Financial Analysis",
    8: "Legal Documentation & Charges",
    9: "Compliance Checklist",
    10: "Appendix",
}

# Cathay Financial green
_GREEN = (0, 112, 60)


async def _require_report(db: AsyncSession, report_id: str) -> Report:
    result = await db.execute(
        select(Report).where(Report.id == report_id, Report.is_deleted == False)
    )
    report = result.scalar_one_or_none()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return report


def _assert_can_view(report: Report, user: User) -> None:
    if user.role in {"admin", "reviewer", "approver"}:
        return
    if report.created_by != user.id:
        raise HTTPException(status_code=403, detail="Access denied")


@router.get("/export/docx")
async def export_docx(
    report_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate and stream a DOCX file for all completed sections."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.warning("export_docx: python-docx not installed, returning 503 for CDN fallback report=%s", report_id)
        raise HTTPException(
            status_code=503,
            detail="python-docx not available on this server — use client-side fallback",
        )

    report = await _require_report(db, report_id)
    _assert_can_view(report, current_user)

    result = await db.execute(
        select(SectionOutput)
        .where(SectionOutput.report_id == report_id, SectionOutput.status == "done")
        .order_by(SectionOutput.section_no)
    )
    outputs = list(result.scalars().all())

    if not outputs:
        logger.warning("export_docx: no completed sections report=%s user=%s", report_id, current_user.id)
        raise HTTPException(status_code=404, detail="No completed sections to export")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Cover title
    title_para = doc.add_heading("Credit Analysis Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_para.runs[0].font.color.rgb = RGBColor(*_GREEN)

    subtitle = report.borrower_name or report_id
    if report.report_type:
        subtitle += f" — {report.report_type}"
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacer

    for output in outputs:
        sec_no = output.section_no
        sec_name = SECTION_NAMES.get(sec_no, f"Section {sec_no}")

        h = doc.add_heading(f"§{sec_no}  {sec_name}", level=1)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(*_GREEN)

        _render_markdown(doc, output.markdown or "", Pt)

        if output != outputs[-1]:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r"[^\w\-]", "_", report.borrower_name or report_id)[:40]
    filename = f"credit_report_{safe_name}.docx"
    logger.info("export_docx: streaming filename=%r sections=%d report=%s user=%s", filename, len(outputs), report_id, current_user.id)

    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── PDF export ────────────────────────────────────────────────────────────────

_PDF_CSS = """\
@page {
  size: A4;
  margin: 2.5cm 2.8cm;
  @top-right { content: "Cathay United Bank — CONFIDENTIAL";
               font-size: 8pt; color: #888; }
  @bottom-center { content: counter(page); font-size: 8pt; color: #888; }
}
body { font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.7; color: #1a1a1a; }
.cover { text-align: center; padding-top: 8cm; page-break-after: always; }
.cover .bank { font-size: 11pt; color: #00703C; font-weight: bold; letter-spacing: 1px; }
.cover .title { font-size: 22pt; font-weight: bold; margin: 24px 0 12px; }
.cover .borrower { font-size: 14pt; color: #333; margin-bottom: 8px; }
.cover .rdate { font-size: 10pt; color: #666; }
h1 { font-size: 13pt; color: #00703C; border-bottom: 2pt solid #00703C;
     padding-bottom: 4px; margin-top: 24px; page-break-after: avoid; }
h2 { font-size: 11pt; color: #005a30; margin-top: 16px; page-break-after: avoid; }
h3 { font-size: 10pt; color: #00703C; }
table { width: 100%; border-collapse: collapse; margin: 10px 0; font-size: 9pt; }
th { background: #e8f5ee; font-weight: bold; padding: 5px 8px;
     border: 1px solid #b2dfce; text-align: left; }
td { padding: 4px 8px; border: 1px solid #d1d5db; vertical-align: top; }
ul, ol { margin: 6px 0 6px 18px; }
li { margin-bottom: 3px; }
p { margin: 6px 0; }
"""


def _md_to_html(md_text: str) -> str:
    """Convert Markdown to HTML, using the `markdown` library if available."""
    try:
        import markdown as md_lib
        return md_lib.markdown(md_text, extensions=["tables", "fenced_code"])
    except ImportError:
        # Basic regex fallback
        t = re.sub(r"^#{1,6} (.+)$", lambda m: f"<h{len(m.group(0).split()[0])}>{m.group(1)}</h{len(m.group(0).split()[0])}>", md_text, flags=re.MULTILINE)
        t = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", t)
        t = re.sub(r"\*(.+?)\*", r"<em>\1</em>", t)
        return "<p>" + t.replace("\n\n", "</p><p>") + "</p>"


@router.get("/export/pdf")
async def export_pdf(
    report_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate and stream a PDF file for all completed sections (requires weasyprint)."""
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="PDF export unavailable — weasyprint is not installed on this server",
        )

    report = await _require_report(db, report_id)
    _assert_can_view(report, current_user)

    result = await db.execute(
        select(SectionOutput)
        .where(SectionOutput.report_id == report_id, SectionOutput.status == "done")
        .order_by(SectionOutput.section_no)
    )
    outputs = list(result.scalars().all())

    if not outputs:
        raise HTTPException(status_code=404, detail="No completed sections to export")

    borrower = report.borrower_name or report_id
    borrower_safe = html_mod.escape(borrower)
    report_date = html_mod.escape(datetime.now().strftime("%d %b %Y"))

    html_parts = [
        f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>Credit Report — {borrower_safe}</title></head><body>
<div class="cover">
  <div class="bank">CATHAY UNITED BANK</div>
  <div class="title">Credit Analysis Report</div>
  <div class="borrower">{borrower_safe}</div>
  <div class="rdate">{report_date}</div>
</div>"""
    ]

    for output in outputs:
        sec_no = output.section_no
        sec_name = SECTION_NAMES.get(sec_no, f"Section {sec_no}")
        html_parts.append(f'<h1>§{sec_no} &nbsp; {sec_name}</h1>')
        html_parts.append(_md_to_html(output.markdown or ""))

    html_parts.append("</body></html>")
    full_html = "\n".join(html_parts)

    pdf_bytes = await asyncio.to_thread(
        lambda: HTML(string=full_html).write_pdf(stylesheets=[CSS(string=_PDF_CSS)])
    )

    safe_name = re.sub(r"[^\w\-]", "_", borrower)[:40]
    filename = f"credit_report_{safe_name}.pdf"
    logger.info(
        "export_pdf: generated filename=%r sections=%d bytes=%d report=%s user=%s",
        filename, len(outputs), len(pdf_bytes), report_id, current_user.id,
    )
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Markdown → python-docx renderer ──────────────────────────────────────────

_TABLE_ROW_RE = re.compile(r"^\|.+\|$")
_TABLE_SEP_RE = re.compile(r"^\|[-:| ]+\|$")
_HEADING_RE = re.compile(r"^(#{2,5})\s+(.+)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.+)$")


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_run_with_inline(para, text: str):
    """Add a run, applying bold/italic inline markdown."""
    # Split on **bold** and *italic* markers and apply formatting
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            if part:
                para.add_run(part)


def _render_table(doc, lines: list[str]):
    data_lines = [l for l in lines if "|" in l and not _TABLE_SEP_RE.match(l)]
    if not data_lines:
        return
    headers = [h.strip() for h in data_lines[0].split("|")[1:-1]]
    rows = []
    for line in data_lines[1:]:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)

    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row — bold + light green shade
    for i, h in enumerate(headers):
        if i < len(table.rows[0].cells):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(_strip_inline(h))
            run.bold = True
            # Light green cell shading
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8F5EE")
            tcPr.append(shd)

    for r_idx, row_vals in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_vals):
            if c_idx < len(row_cells):
                row_cells[c_idx].text = _strip_inline(val)

    doc.add_paragraph()  # spacing after table


def _render_markdown(doc, markdown: str, Pt):
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if m := _HEADING_RE.match(line):
            level = min(len(m.group(1)), 4)
            doc.add_heading(_strip_inline(m.group(2)), level=level)

        elif _TABLE_SEP_RE.match(line) or _TABLE_ROW_RE.match(line):
            table_lines = []
            while i < len(lines) and (
                _TABLE_ROW_RE.match(lines[i]) or _TABLE_SEP_RE.match(lines[i])
            ):
                table_lines.append(lines[i])
                i += 1
            _render_table(doc, table_lines)
            continue

        elif m := _BULLET_RE.match(line):
            p = doc.add_paragraph(style="List Bullet")
            _add_run_with_inline(p, m.group(1))

        elif m := _NUMBERED_RE.match(line):
            p = doc.add_paragraph(style="List Number")
            _add_run_with_inline(p, m.group(1))

        elif not line.strip():
            pass  # blank lines = natural paragraph separation

        else:
            stripped = _strip_inline(line)
            if stripped:
                p = doc.add_paragraph()
                _add_run_with_inline(p, line)  # keep inline markup

        i += 1
